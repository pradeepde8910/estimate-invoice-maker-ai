"""
File handling utilities — supports PDF, DOCX, plain text, and URL ingestion.
"""

import ipaddress
import os
import re
import socket
import tempfile
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

import requests

_ALLOWED_URL_SCHEMES = {"http", "https"}


class UnsafeURLError(ValueError):
    """Raised when a user-supplied URL points somewhere download_url() refuses to fetch."""


def _is_public_ip(ip: str) -> bool:
    addr = ipaddress.ip_address(ip)
    return not (
        addr.is_private
        or addr.is_loopback
        or addr.is_link_local
        or addr.is_multicast
        or addr.is_reserved
        or addr.is_unspecified
    )


def _assert_safe_url(url: str) -> None:
    """Blocks SSRF: rejects non-http(s) schemes and any hostname that
    resolves to a private/loopback/link-local address (cloud metadata
    endpoints, localhost, internal services)."""
    parsed = urlparse(url)
    if parsed.scheme not in _ALLOWED_URL_SCHEMES:
        raise UnsafeURLError(f"URL scheme '{parsed.scheme}' is not allowed; only http/https are permitted.")

    hostname = parsed.hostname
    if not hostname:
        raise UnsafeURLError("URL has no hostname.")

    try:
        resolved = socket.getaddrinfo(hostname, None)
    except socket.gaierror as e:
        raise UnsafeURLError(f"Could not resolve host '{hostname}': {e}")

    for _, _, _, _, sockaddr in resolved:
        ip = sockaddr[0]
        if not _is_public_ip(ip):
            raise UnsafeURLError(
                f"'{hostname}' resolves to a non-public address ({ip}). "
                "Internal/private/loopback/link-local addresses are blocked."
            )


def read_file(file_path: str) -> tuple[str, str]:
    """
    Read a file and return (content_or_path, file_type).
    For PDFs, returns the file path (OCR handles them).
    For DOCX, extracts text.
    For text files, reads content directly.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if ext == ".pdf":
        return str(path.resolve()), "pdf"

    elif ext == ".docx":
        return _read_docx(file_path), "docx"

    elif ext in (".txt", ".md", ".rst", ".csv"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read(), "text"

    else:
        # Try reading as text
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read(), "text"
        except Exception:
            raise ValueError(f"Unsupported file type: {ext}")


def _read_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        tables_text = []
        for table in doc.tables:
            seen_cells = set()
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell in seen_cells:
                        continue
                    seen_cells.add(cell)
                    try:
                        row_data.append(cell.text.strip())
                    except Exception:
                        row_data.append("")
                if any(row_data):
                    tables_text.append(" | ".join(row_data))
            tables_text.append("")  # separator between tables

        all_text = "\n".join(paragraphs)
        if tables_text:
            all_text += "\n\n--- TABLES ---\n" + "\n".join(tables_text)
        return all_text
    except Exception as docx_error:
        # Fallback to basic XML parsing if python-docx fails
        import zipfile
        import xml.etree.ElementTree as ET
        try:
            with zipfile.ZipFile(file_path) as z:
                xml_content = z.read('word/document.xml')
            tree = ET.XML(xml_content)
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            texts = tree.findall('.//w:t', namespaces)
            return '\n'.join([t.text for t in texts if t.text])
        except Exception as fallback_error:
            raise ValueError(f"Failed to read DOCX file. It may be corrupted or not a valid .docx document. (Details: {docx_error})")


MAX_DOWNLOAD_BYTES = 50 * 1024 * 1024  # 50 MB


def _write_capped(response, tmp) -> None:
    written = 0
    for chunk in response.iter_content(chunk_size=8192):
        written += len(chunk)
        if written > MAX_DOWNLOAD_BYTES:
            tmp.close()
            raise UnsafeURLError(f"Response exceeded the {MAX_DOWNLOAD_BYTES // (1024 * 1024)}MB download limit.")
        tmp.write(chunk)


def download_url(url: str) -> tuple[str, str]:
    """
    Download content from a URL.
    If it's a PDF, save to temp file and return path.
    Otherwise, return the text content.
    """
    _assert_safe_url(url)

    # No redirects: an open redirect on an otherwise-public host is a
    # well-known way to bypass host allow-listing and reach an internal
    # address anyway. Failing closed here is simpler and safer than
    # re-validating every hop.
    response = requests.get(url, timeout=30, stream=True, allow_redirects=False)
    if response.is_redirect or response.status_code in (301, 302, 303, 307, 308):
        raise UnsafeURLError("The URL responded with a redirect, which is not followed (SSRF protection).")
    response.raise_for_status()

    content_length = response.headers.get("Content-Length")
    if content_length and int(content_length) > MAX_DOWNLOAD_BYTES:
        raise UnsafeURLError(f"Response exceeded the {MAX_DOWNLOAD_BYTES // (1024 * 1024)}MB download limit.")

    content_type = response.headers.get("Content-Type", "").lower()

    if "pdf" in content_type or url.lower().endswith(".pdf"):
        # Save PDF to temp file
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        _write_capped(response, tmp)
        tmp.close()
        return tmp.name, "pdf"

    elif "word" in content_type or url.lower().endswith(".docx"):
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        _write_capped(response, tmp)
        tmp.close()
        return _read_docx(tmp.name), "docx"

    else:
        # Treat as text/HTML
        return response.text, "text"


def is_url(input_str: str) -> bool:
    """Check if the input string is a URL."""
    return bool(re.match(r"https?://", input_str.strip()))


def detect_input_type(input_str: str) -> str:
    """Detect whether the input is a file path, URL, or raw text."""
    stripped = input_str.strip()
    if is_url(stripped):
        return "url"
    elif os.path.isfile(stripped):
        return "file"
    else:
        return "text"
