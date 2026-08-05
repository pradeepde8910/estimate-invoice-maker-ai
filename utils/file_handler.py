"""
File handling utilities — supports PDF, DOCX, plain text, and URL ingestion.
"""

import os
import re
import tempfile
from pathlib import Path
from typing import Optional

import requests


def read_file(file_path: str) -> tuple[str, str]:
    """
    Read a file and return (content_or_path, file_type).
    For PDFs, returns the file path (OCR handles them).
    For DOCX, extracts text.
    For text files, reads content directly.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if ext == ".pdf":
        return str(path.resolve()), "pdf"

    elif ext == ".docx":
        return _read_docx(file_path), "docx"

    elif ext in (".txt", ".md", ".rst", ".csv"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read(), "text"

    else:
        # Try reading as text
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read(), "text"
        except Exception:
            raise ValueError(f"Unsupported file type: {ext}")


def _read_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            tables_text.append(" | ".join(row_data))
        tables_text.append("")  # separator between tables

    all_text = "\n".join(paragraphs)
    if tables_text:
        all_text += "\n\n--- TABLES ---\n" + "\n".join(tables_text)
    return all_text


def download_url(url: str) -> tuple[str, str]:
    """
    Download content from a URL.
    If it's a PDF, save to temp file and return path.
    Otherwise, return the text content.
    """
    response = requests.get(url, timeout=30, stream=True)
    response.raise_for_status()

    content_type = response.headers.get("Content-Type", "").lower()

    if "pdf" in content_type or url.lower().endswith(".pdf"):
        # Save PDF to temp file
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        for chunk in response.iter_content(chunk_size=8192):
            tmp.write(chunk)
        tmp.close()
        return tmp.name, "pdf"

    elif "word" in content_type or url.lower().endswith(".docx"):
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        for chunk in response.iter_content(chunk_size=8192):
            tmp.write(chunk)
        tmp.close()
        return _read_docx(tmp.name), "docx"

    else:
        # Treat as text/HTML
        return response.text, "text"


def is_url(input_str: str) -> bool:
    """Check if the input string is a URL."""
    return bool(re.match(r"https?://", input_str.strip()))


def detect_input_type(input_str: str) -> str:
    """Detect whether the input is a file path, URL, or raw text."""
    stripped = input_str.strip()
    if is_url(stripped):
        return "url"
    elif os.path.isfile(stripped):
        return "file"
    else:
        return "text"
